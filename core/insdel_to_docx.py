"""
Markdown (<ins>/<del> + [^N]) → DOCX with native Track Changes + Footnotes

將 Markdown 中的：
  <ins> / <del> → Word 原生追蹤修訂（插入藍字底線 / 刪除紅字刪除線）
  [^N] / [^N]: → Word 原生腳注（頁底附註，非尾注）

不使用 Pandoc，不使用 Word Compare。
直接透過 python-docx 操作 XML 層。
"""

import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from utils.file_utils import resolve_output_conflict


WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


class InsDelToDocx:
    """解析 <ins>/<del> + [^N] 標記，輸出 Word 原生追蹤修訂 + 腳注 DOCX"""

    def __init__(self, author: str = "AI Revision"):
        self.author = author
        self._revision_id = 0
        self._fn_counter = 0          # footnote ID counter
        self._fn_map = {}              # fn_label → (fn_id, fn_text)
        self._next_fn_id = 1           # Word footnote internal ID
        self._fn_part = None           # footnotes XML part
        self._fn_rels = None           # relationships for footnotes part

    def convert(self, markdown_text: str, output_path: str) -> tuple[bool, str, str | None]:
        try:
            final_output_path, renamed = resolve_output_conflict(output_path)
            doc = Document()

            # 啟用追蹤修訂
            self._enable_track_changes(doc)

            # 準備腳注基礎設施
            self._prepare_footnotes(doc)

            # 設定預設字型
            style = doc.styles["Normal"]
            style.font.size = Pt(12)
            style.font.name = "Microsoft JhengHei"

            # ── 第一遍：擷取腳注定義 [^N]: ... ──
            lines = markdown_text.split("\n")
            body_lines = []
            for line in lines:
                m = re.match(r"^\[\^(\d+)\]:\s*(.*)", line.rstrip())
                if m:
                    label = m.group(1)
                    text = m.group(2).strip()
                    self._fn_map[label] = text
                else:
                    body_lines.append(line)

            # ── 第二遍：渲染正文 ──
            for line in body_lines:
                line = line.rstrip()
                if not line.strip():
                    doc.add_paragraph("")
                    continue

                if "<ins>" in line or "<del>" in line or re.search(r"\[\^\d+\]", line):
                    self._add_tracked_paragraph_with_fn(doc, line)
                elif line.startswith("#"):
                    self._add_markdown_heading(doc, line)
                else:
                    self._add_plain_paragraph(doc, line)

            doc.save(final_output_path)
            if renamed:
                return True, f"追蹤修訂 DOCX 已產出；原輸出檔被占用，已改存為：{final_output_path}", final_output_path
            return True, f"追蹤修訂 DOCX 已產出：{final_output_path}", final_output_path

        except Exception as e:
            import traceback
            traceback.print_exc()
            return False, f"轉換失敗：{e}", None

    # ═══════════════════════════════════════════════════════════
    #  追蹤修訂基礎
    # ═══════════════════════════════════════════════════════════

    def _enable_track_changes(self, doc):
        """文件層級啟用 trackRevisions"""
        settings_elem = doc.settings.element
        TR = qn("w:trackRevisions")
        if settings_elem.find(TR) is None:
            settings_elem.append(OxmlElement("w:trackRevisions"))

    # ═══════════════════════════════════════════════════════════
    #  腳注基礎設施（關鍵：區分腳注與尾注）
    # ═══════════════════════════════════════════════════════════

    def _prepare_footnotes(self, doc):
        """
        初始化 Word 腳注（footnotes）part。
        使用 python-docx 內部機制確保類型正確。
        """
        from docx.opc.part import XmlPart
        from docx.opc.packuri import PackURI

        footnotes_uri = PackURI("/word/footnotes.xml")
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"

        # 從 lxml 解析基礎 XML
        fn_element = etree.fromstring(self._build_footnotes_xml())

        # 直接以 Part.__init__ 的參數順序建立
        self._fn_part = XmlPart(footnotes_uri, content_type, fn_element, doc.part.package)

        # 建立主文檔到 footnotes 的關係
        doc.part.relate_to(
            self._fn_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        )

    def _build_footnotes_xml(self) -> bytes:
        """建立 footnotes.xml 的最小基礎結構"""
        xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )
        return xml.encode("utf-8")

    def _add_footnote_definition(self, fn_text: str) -> int:
        """
        在 footnotes.xml 中加入一個腳注定義。
        回傳 footnote ID（供 reference 引用）。

        腳注格式：10pt 字、單行間距、凸排對齊。
        """
        fn_id = self._next_fn_id
        self._next_fn_id += 1

        # 建立 footnote XML
        footnote = OxmlElement("w:footnote")
        footnote.set(qn("w:id"), str(fn_id))

        # ── 段落（含格式）──
        p = OxmlElement("w:p")

        # 段落屬性：單行間距 + 首行懸掛 0.5 字符
        pPr = OxmlElement("w:pPr")
        # 行距：240 = 單行間距
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:before"), "0")
        pPr.append(spacing)
        # 首行懸掛 0.5 字符：w:left + w:hanging 同值 = 後續行內縮，首行不縮
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "120")
        ind.set(qn("w:hanging"), "120")
        pPr.append(ind)
        p.append(pPr)

        # 腳注編號 run（10pt + 上標）
        r1 = OxmlElement("w:r")
        rPr1 = OxmlElement("w:rPr")
        sz1 = OxmlElement("w:sz")
        sz1.set(qn("w:val"), "20")          # 10pt = 20 half-points
        rPr1.append(sz1)
        szCs1 = OxmlElement("w:szCs")
        szCs1.set(qn("w:val"), "20")
        rPr1.append(szCs1)
        vertAlign = OxmlElement("w:vertAlign")
        vertAlign.set(qn("w:val"), "superscript")
        rPr1.append(vertAlign)
        r1.append(rPr1)
        fnRef = OxmlElement("w:footnoteRef")
        r1.append(fnRef)
        p.append(r1)

        # 腳注文字 run（10pt）
        r2 = OxmlElement("w:r")
        rPr2 = OxmlElement("w:rPr")
        sz2 = OxmlElement("w:sz")
        sz2.set(qn("w:val"), "20")          # 10pt
        rPr2.append(sz2)
        szCs2 = OxmlElement("w:szCs")
        szCs2.set(qn("w:val"), "20")
        rPr2.append(szCs2)
        r2.append(rPr2)
        t2 = OxmlElement("w:t")
        t2.set(qn("xml:space"), "preserve")
        t2.text = " " + fn_text
        r2.append(t2)
        p.append(r2)

        footnote.append(p)

        # 附加到 footnotes part 的 XML 根部
        fn_root = self._fn_part._element
        fn_root.append(footnote)

        return fn_id

    def _insert_footnote_reference(self, paragraph, fn_id: int):
        """
        在段落末尾插入一個腳注引用 (w:footnoteReference)。
        """
        # 建立引用 run
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        vertAlign = OxmlElement("w:vertAlign")
        vertAlign.set(qn("w:val"), "superscript")
        rPr.append(vertAlign)
        r.append(rPr)

        ref = OxmlElement("w:footnoteReference")
        ref.set(qn("w:id"), str(fn_id))
        r.append(ref)

        paragraph._element.append(r)

    # ═══════════════════════════════════════════════════════════
    #  段落渲染
    # ═══════════════════════════════════════════════════════════

    def _add_plain_paragraph(self, doc, text: str):
        p = doc.add_paragraph()

        # 檢查是否有腳注引用
        fn_refs = list(re.finditer(r"\[\^(\d+)\]", text))
        if fn_refs:
            # 分割文字和引用
            last_end = 0
            for m in fn_refs:
                # 前置文字
                if m.start() > last_end:
                    run = p.add_run(text[last_end:m.start()])
                    run.font.size = Pt(12)
                # 腳注引用
                label = m.group(1)
                fn_text = self._fn_map.get(label, "")
                fn_id = self._add_footnote_definition(fn_text)
                self._insert_footnote_reference(p, fn_id)
                last_end = m.end()
            # 剩餘文字
            if last_end < len(text):
                run = p.add_run(text[last_end:])
                run.font.size = Pt(12)
        else:
            run = p.add_run(text)
            run.font.size = Pt(12)

    def _add_markdown_heading(self, doc, line: str):
        match = re.match(r"^(#+)\s+(.+)", line)
        if match:
            level = min(len(match.group(1)), 9)
            text = match.group(2)
            doc.add_heading(text, level=level)

    def _add_tracked_paragraph_with_fn(self, doc, line: str):
        """處理同時含 <ins>/<del> 和 [^N] 的行"""
        p = doc.add_paragraph()

        # 正則：分割 <del>...</del>、<ins>...</ins>、[^N]
        pattern = r"(<del>.*?</del>|<ins>.*?</ins>|\[\^\d+\])"
        parts = re.split(pattern, line)

        for part in parts:
            if not part:
                continue

            # <del> 區塊
            m_del = re.match(r"^<del>(.*?)</del>$", part, re.DOTALL)
            if m_del:
                text = m_del.group(1)
                if text:
                    self._add_deletion_run(p, text)
                continue

            # <ins> 區塊
            m_ins = re.match(r"^<ins>(.*?)</ins>$", part, re.DOTALL)
            if m_ins:
                text = m_ins.group(1)
                if text:
                    self._add_insertion_run(p, text)
                continue

            # [^N] 腳注引用
            m_fn = re.match(r"^\[\^(\d+)\]$", part)
            if m_fn:
                label = m_fn.group(1)
                fn_text = self._fn_map.get(label, "")
                fn_id = self._add_footnote_definition(fn_text)
                self._insert_footnote_reference(p, fn_id)
                continue

            # 普通文字
            if part.strip():
                run = p.add_run(part)
                run.font.size = Pt(12)

    # ═══════════════════════════════════════════════════════════
    #  追蹤修訂 XML helpers
    # ═══════════════════════════════════════════════════════════

    def _next_rev_id(self) -> int:
        self._revision_id += 1
        return self._revision_id

    def _add_deletion_run(self, paragraph, text: str):
        rev_id = self._next_rev_id()
        now = datetime.utcnow().isoformat() + "Z"

        del_elem = OxmlElement("w:del")
        del_elem.set(qn("w:id"), str(rev_id))
        del_elem.set(qn("w:author"), self.author)
        del_elem.set(qn("w:date"), now)

        # 只標記「這是刪除修訂」，不要把紅色/刪除線硬寫進文字本身。
        # 顯示樣式交給 Word 的修訂顯示設定，這樣拒絕刪除後才不會殘留格式。
        del_run = OxmlElement("w:r")

        del_text = OxmlElement("w:delText")
        del_text.set(qn("xml:space"), "preserve")
        del_text.text = text
        del_run.append(del_text)

        del_elem.append(del_run)
        paragraph._element.append(del_elem)

    def _add_insertion_run(self, paragraph, text: str):
        rev_id = self._next_rev_id()
        now = datetime.utcnow().isoformat() + "Z"

        ins_elem = OxmlElement("w:ins")
        ins_elem.set(qn("w:id"), str(rev_id))
        ins_elem.set(qn("w:author"), self.author)
        ins_elem.set(qn("w:date"), now)

        # 同樣不要把藍字底線寫進新增內容本身。
        # 接受修訂後，新增文字就會自然回到正文格式。
        ins_run = OxmlElement("w:r")

        t_elem = OxmlElement("w:t")
        t_elem.set(qn("xml:space"), "preserve")
        t_elem.text = text
        ins_run.append(t_elem)

        ins_elem.append(ins_run)
        paragraph._element.append(ins_elem)
